"""
Reorganize documents according to the correct classification.
"""

import os
import shutil

# Create folder structure
base_dir = "/Users/craig/Desktop/Legal-amendment"
output_dir = os.path.join(base_dir, "Final_Word_Documents")

folders = {
    "RP1_Old": os.path.join(output_dir, "RP1_Old"),
    "RP1_UCT_New": os.path.join(output_dir, "RP1_UCT_New"),
    "RP1_WHC_New": os.path.join(output_dir, "RP1_WHC_New"),
    "RP2_WHC": os.path.join(output_dir, "RP2_WHC")
}

# Create the directories
for folder in folders.values():
    os.makedirs(folder, exist_ok=True)
    print(f"Created directory: {folder}")

# Define the source files and their destinations
files_to_copy = [
    # Old RP1 Document
    {
        "source": os.path.join(base_dir, "OLD_DTAs", "WHC_DTA_RP1_0911 (1).docx"),
        "dest": os.path.join(folders["RP1_Old"], "RP1_WHC_Original.docx"),
        "description": "Original RP1 WHC Document"
    },
    
    # New WHC RP1 Document
    {
        "source": os.path.join(base_dir, "OLD_DTAs", "DTA_HEAT001_WHC TEMPLATE_20240131_comments_LvA_CPedits_MF1-1.docx"),
        "dest": os.path.join(folders["RP1_WHC_New"], "RP1_WHC_Original.docx"),
        "description": "New WHC RP1 Document (Original)"
    },
    
    # New UCT RP1 Document
    {
        "source": os.path.join(base_dir, "OLD_DTAs", "DTA_Template_HEAT001_v2.0_UCT_DRAFT 10.09.2024_TC.docx"),
        "dest": os.path.join(folders["RP1_UCT_New"], "RP1_UCT_Original.docx"),
        "description": "New UCT RP1 Document (Original)"
    },
    
    # RP2 Document
    {
        "source": os.path.join(base_dir, "OLD_DTAs", "Final_Formatted_DTA_HEAT002_WHC_RP2_0903 (3).docx"),
        "dest": os.path.join(folders["RP2_WHC"], "RP2_WHC_Original.docx"),
        "description": "RP2 WHC Document (Original)"
    }
]

# Copy the files to their destinations
for file_info in files_to_copy:
    source = file_info["source"]
    dest = file_info["dest"]
    description = file_info["description"]
    
    if os.path.exists(source):
        shutil.copy2(source, dest)
        print(f"Copied {description} from {source} to {dest}")
    else:
        print(f"Warning: {description} not found at {source}")

# Now let's create the tracked changes versions
from docx import Document
from docx.shared import RGBColor

def add_tracked_changes(input_docx, output_docx, changes):
    """Add tracked changes to a Word document."""
    try:
        doc = Document(input_docx)
        print(f"Successfully opened {input_docx}")
    except Exception as e:
        print(f"Error opening document: {e}")
        return False
    
    # Process all paragraphs
    for change in changes:
        change_type = change.get('type')
        search_text = change.get('search_text')
        replacement = change.get('replacement', '')
        
        # Find and process the text in the document
        for paragraph in doc.paragraphs:
            if search_text in paragraph.text:
                # Create a new paragraph with the changes
                text_parts = paragraph.text.split(search_text)
                paragraph.clear()
                
                # Add text before the match
                if text_parts[0]:
                    run = paragraph.add_run(text_parts[0])
                
                # Add the changed text with appropriate styling
                if change_type == 'deletion':
                    run = paragraph.add_run(search_text)
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                    run.font.strike = True
                elif change_type == 'addition':
                    if search_text:  # If we're replacing something
                        run = paragraph.add_run(search_text)
                        run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                        run.font.strike = True
                    
                    run = paragraph.add_run(replacement)
                    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
                
                # Add text after the match
                if len(text_parts) > 1:
                    run = paragraph.add_run(text_parts[1])
    
    # Save the modified document
    doc.save(output_docx)
    print(f"Saved modifications to {output_docx}")
    return True

# Define the changes for our documents
changes = [
    # Common changes for all documents
    {
        'type': 'addition',
        'search_text': '1.20 "Consortium Shared Data" means',
        'replacement': '\n\n1.21 "Azure Cloud Platform" means the Microsoft Azure cloud computing service that will serve as the HE²AT Center Primary Repository.\n\n1.22 "Cloud Migration" means the process of transferring the Original Study Data and any derived data sets from the existing on-premises infrastructure to the Azure Cloud Platform.\n\n1.23 "Data Access Committee or DAC" means the committee established by the HE²AT Center to review and approve data access requests from external researchers according to established criteria and protocols, which shall continue to function after the conclusion of the HE²AT Center Project.\n\n1.24 "Data Access Levels" means the tiered access system implemented by the HE²AT Center consisting of:\n• Level 0: Original Study Data - Raw, unprocessed data with restricted access to Core Data Team only\n• Level 1: Consortium Shared Data - Processed data shared only among HE²AT Center Consortium partners\n• Level 2: De-identified Data - Retained by HE²AT Center for approved external researcher access\n• Level 3: Inferential Data - Aggregated and anonymized data available for open access\n\n1.25 "External Researcher" means any qualified researcher who is not a member of the HE²AT Center Consortium but who has been approved by the Data Access Committee to access Level 2 data for specific research purposes.\n\n1.26 "Post-Project Data Use" means the continued storage, access, and use of the data after the conclusion of the HE²AT Center Project in accordance with this Amendment.\n\n1.27 "Successor Governance Entity" means any entity or institution that assumes responsibility for the governance, maintenance, and oversight of the Post-Project Data Repository after the conclusion of the HE²AT Center Project.'
    },
    {
        'type': 'deletion',
        'search_text': 'This Agreement shall commence on the Commencement Date and shall terminate on completion of the HE²AT Project',
        'replacement': 'This Agreement shall commence on the Commencement Date and shall remain in effect in perpetuity with respect to the Post-Project Data Use provisions set forth in Section 2.12 of this Amendment, unless terminated earlier in accordance with the provisions of this Agreement. The Data Provider specifically acknowledges and agrees that the Post-Project Data Use provisions shall survive the termination of the HE²AT Center Project'
    },
    {
        'type': 'addition',
        'search_text': '2.4 Data Provider retains ownership of the Original Study Data and retains all rights to distribute the Original Study Data to other third parties.',
        'replacement': '\n\nThe Data Provider hereby grants the Data Recipient a perpetual, irrevocable, worldwide, non-exclusive license to store the Original Study Data in the Azure Cloud Platform as part of the HE²AT Center\'s Cloud Migration.\n\nThe Data Provider hereby authorizes the Data Recipient to process and transform the Original Study Data to create derived data sets at Levels 1, 2, and 3.\n\nThe Data Provider hereby authorizes the Data Recipient to retain the derived data sets for Post-Project Data Use as specifically authorized in this Amendment.\n\nThe Data Provider hereby authorizes the Data Recipient to grant access to Level 2 data to External Researchers in accordance with the procedures set forth in this Amendment.\n\nThis expanded license does not transfer ownership of the Original Study Data, which remains with the Data Provider.'
    },
    {
        'type': 'deletion',
        'search_text': 'The Data Provider acknowledges and agrees that initially the Original Study Data shall be accessible only to the Core Team for purposes of',
        'replacement': 'The Data Provider acknowledges and agrees that initially the Original Study Data shall be accessible only to the Core HE²AT Center Data Management Team for purposes of'
    },
    {
        'type': 'addition',
        'search_text': 'as set out in the HE²AT Center Data Management Plan.',
        'replacement': ' as set out in the HE²AT Center Data Management Plan. Following the Cloud Migration, the Original Study Data will be classified as Level 0 data in the Azure Cloud Platform\'s tiered data access system and will remain accessible only to the Core Data Team.'
    },
    {
        'type': 'addition',
        'search_text': '2.18 The Data Provider',
        'replacement': '\n\n2.19 Cloud Storage Infrastructure and Data Migration Authorization\n\n2.19.1 The Data Provider hereby irrevocably authorizes the Data Recipient to:\n(a) migrate the Original Study Data from on-premises infrastructure to the Azure Cloud Platform;\n(b) store and process the Original Study Data and all derived data sets in the Azure Cloud Platform; and\n(c) implement the tiered Data Access Levels system described in this Amendment.\n\n2.20 External Researcher Access Authorization\n\n2.20.1 The Data Provider hereby explicitly and irrevocably authorizes and grants permission for appropriately de-identified data derived from the Original Study Data (Level 2 Data) to be made available to External Researchers who are not members of the HE²AT Center Consortium, subject to the following conditions:\n(a) All External Researcher access requests must be reviewed and approved by the Data Access Committee.\n(b) External Researchers must sign a legally binding Data Use Agreement.\n(c) External Researchers must commit to appropriate citation of both the HE²AT Center and the original data sources.\n(d) External Researchers will only have access to Level 2 data (de-identified).\n(e) All External Researcher access will be monitored and logged.\n(f) The Data Access Committee shall maintain the right to revoke access for any External Researcher.\n\n2.21 Post-Project Data Use and Long-Term Data Retention Authorization\n\n2.21.1 The Data Provider hereby explicitly and irrevocably authorizes and grants permission that the data derived from the Original Study Data shall be retained and may continue to be used beyond the conclusion of the HE²AT Center Project as follows:\n(a) Level 0 Data (Original Study Data): Shall be retained for a period of 5 (five) years.\n(b) Level 1 Data (Consortium Shared Data): Shall be retained for a period of 10 (ten) years.\n(c) Level 2 Data (De-identified Data): Shall be retained indefinitely as a scientific resource.\n(d) Level 3 Data (Inferential Data): Shall be retained indefinitely as an open scientific resource.\n\n2.18 The Data Provider'
    },
    {
        'type': 'addition',
        'search_text': '12.5 The provisions of this Agreement that by their nature are intended to survive termination or expiration of the Agreement shall survive such termination or expiration and shall remain in full force and effect.',
        'replacement': '\n\n12.6 The Parties expressly acknowledge and agree that the provisions of Sections 2.20 and 2.21 of this Amendment regarding External Researcher Access and Post-Project Data Use shall survive the termination of the Agreement and the conclusion of the HE²AT Center Project.'
    },
    {
        'type': 'deletion',
        'search_text': '12.6 The Data Recipient hereby acknowledges',
        'replacement': '12.7 The Data Recipient hereby acknowledges'
    },
    {
        'type': 'addition',
        'search_text': 'IN WITNESS WHEREOF, the Parties have executed this Agreement as of the Effective Date.',
        'replacement': 'IN WITNESS WHEREOF, the Parties have executed this Agreement as of the Amendment Effective Date.'
    }
]

# Create tracked changes versions for each document
tracked_changes_files = [
    {
        "input": os.path.join(folders["RP1_WHC_New"], "RP1_WHC_Original.docx"),
        "output": os.path.join(folders["RP1_WHC_New"], "RP1_WHC_Tracked_Changes.docx"),
        "description": "New WHC RP1 Document with Tracked Changes"
    },
    {
        "input": os.path.join(folders["RP1_UCT_New"], "RP1_UCT_Original.docx"),
        "output": os.path.join(folders["RP1_UCT_New"], "RP1_UCT_Tracked_Changes.docx"),
        "description": "New UCT RP1 Document with Tracked Changes"
    },
    {
        "input": os.path.join(folders["RP2_WHC"], "RP2_WHC_Original.docx"),
        "output": os.path.join(folders["RP2_WHC"], "RP2_WHC_Tracked_Changes.docx"),
        "description": "RP2 WHC Document with Tracked Changes"
    }
]

# Process each file with tracked changes
for file_info in tracked_changes_files:
    input_file = file_info["input"]
    output_file = file_info["output"]
    description = file_info["description"]
    
    if os.path.exists(input_file):
        result = add_tracked_changes(input_file, output_file, changes)
        if result:
            print(f"Successfully created {description} at {output_file}")
        else:
            print(f"Failed to create {description}")
    else:
        print(f"Warning: Input file for {description} not found at {input_file}")

print("\nDocument reorganization complete!")
print(f"All files are now organized in: {output_dir}")
print("\nFile structure:")
for folder_name, folder_path in folders.items():
    print(f"\n{folder_name}:")
    for file in os.listdir(folder_path):
        print(f"  - {file}")
