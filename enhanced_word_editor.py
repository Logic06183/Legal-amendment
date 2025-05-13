"""
Enhanced script to edit Word documents with tracked changes for all versions.
This script organizes documents into appropriate folders and adds red for deletions and blue for additions.
"""

import os
import shutil
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_tracked_changes(input_docx, output_docx, changes):
    """
    Add tracked changes to a Word document.
    
    Args:
        input_docx: Path to the input Word document
        output_docx: Path to save the modified document
        changes: List of dictionaries with changes to make
    """
    try:
        doc = Document(input_docx)
        print(f"Successfully opened {input_docx}")
    except Exception as e:
        print(f"Error opening document: {e}")
        return False
    
    # Process all paragraphs
    for change in changes:
        change_type = change.get('type')
        search_text = change.get('search_text')
        replacement = change.get('replacement', '')
        
        # Find and process the text in the document
        for paragraph in doc.paragraphs:
            if search_text in paragraph.text:
                # Create a new paragraph with the changes
                text_parts = paragraph.text.split(search_text)
                paragraph.clear()
                
                # Add text before the match
                if text_parts[0]:
                    run = paragraph.add_run(text_parts[0])
                
                # Add the changed text with appropriate styling
                if change_type == 'deletion':
                    run = paragraph.add_run(search_text)
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                    run.font.strike = True
                elif change_type == 'addition':
                    if search_text:  # If we're replacing something
                        run = paragraph.add_run(search_text)
                        run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                        run.font.strike = True
                    
                    run = paragraph.add_run(replacement)
                    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
                
                # Add text after the match
                if len(text_parts) > 1:
                    run = paragraph.add_run(text_parts[1])
    
    # Save the modified document
    doc.save(output_docx)
    print(f"Saved modifications to {output_docx}")
    return True

def copy_to_destination(source_path, dest_path):
    """Copy a file to the destination, creating directories if needed"""
    os.makedirs(os.path.dirname(dest_path), exist_ok=True)
    shutil.copy2(source_path, dest_path)
    print(f"Copied {source_path} to {dest_path}")

def process_rp2_document():
    """Process the RP2 document with tracked changes (WHC only)"""
    # Define base paths
    input_path = os.path.join("OLD_DTAs_Organized", "RP2_DTA_WHC_Main.docx")
    output_dir = os.path.join("Word_Documents_With_Tracked_Changes", "RP2_WHC")
    output_path = os.path.join(output_dir, "RP2_DTA_WHC_Tracked.docx")
    
    # Ensure the output directory exists
    os.makedirs(output_dir, exist_ok=True)
    
    # Define the changes for RP2
    changes = [
        # Section 1 changes - Definitions
        {
            'type': 'addition',
            'search_text': '1.20 "Consortium Shared Data" means',
            'replacement': '\n\n1.21 "Azure Cloud Platform" means the Microsoft Azure cloud computing service that will serve as the HE²AT Center Primary Repository.\n\n1.22 "Cloud Migration" means the process of transferring the Original Study Data and any derived data sets from the existing on-premises infrastructure to the Azure Cloud Platform.\n\n1.23 "Data Access Committee or DAC" means the committee established by the HE²AT Center to review and approve data access requests from external researchers according to established criteria and protocols, which shall continue to function after the conclusion of the HE²AT Center Project.\n\n1.24 "Data Access Levels" means the tiered access system implemented by the HE²AT Center consisting of:\n• Level 0: Original Study Data - Raw, unprocessed data with restricted access to Core Data Team only\n• Level 1: Consortium Shared Data - Processed data shared only among HE²AT Center Consortium partners\n• Level 2: RP2 De-identified Data - Retained by HE²AT Center for approved external researcher access\n• Level 3: Inferential Data - Aggregated and anonymized data available for open access\n\n1.25 "External Researcher" means any qualified researcher who is not a member of the HE²AT Center Consortium but who has been approved by the Data Access Committee to access Level 2 data for specific research purposes.\n\n1.26 "Post-Project Data Use" means the continued storage, access, and use of the data after the conclusion of the HE²AT Center Project in accordance with this Amendment.\n\n1.27 "Successor Governance Entity" means any entity or institution that assumes responsibility for the governance, maintenance, and oversight of the Post-Project Data Repository after the conclusion of the HE²AT Center Project.'
        },
        # Section 2.1 changes
        {
            'type': 'deletion',
            'search_text': 'This Agreement shall commence on the Commencement Date and shall terminate on completion of the HE²AT Project',
            'replacement': 'This Agreement shall commence on the Commencement Date and shall remain in effect in perpetuity with respect to the Post-Project Data Use provisions set forth in Section 2.12 of this Amendment, unless terminated earlier in accordance with the provisions of this Agreement. The Data Provider specifically acknowledges and agrees that the Post-Project Data Use provisions shall survive the termination of the HE²AT Center Project'
        },
        # Add section after 2.4
        {
            'type': 'addition',
            'search_text': '2.4 Data Provider retains ownership of the Original Study Data and retains all rights to distribute the Original Study Data to other third parties.',
            'replacement': '\n\nThe Data Provider hereby grants the Data Recipient a perpetual, irrevocable, worldwide, non-exclusive license to store the Original Study Data in the Azure Cloud Platform as part of the HE²AT Center\'s Cloud Migration.\n\nThe Data Provider hereby authorizes the Data Recipient to process and transform the Original Study Data to create derived data sets at Levels 1, 2, and 3.\n\nThe Data Provider hereby authorizes the Data Recipient to retain the derived data sets for Post-Project Data Use as specifically authorized in this Amendment.\n\nThe Data Provider hereby authorizes the Data Recipient to grant access to Level 2 data to External Researchers in accordance with the procedures set forth in this Amendment.\n\nThis expanded license does not transfer ownership of the Original Study Data, which remains with the Data Provider.'
        },
        # New sections 2.19-2.21
        {
            'type': 'addition',
            'search_text': '2.18 The Data Provider',
            'replacement': '\n\n2.19 Cloud Storage Infrastructure and Data Migration Authorization\n\n2.19.1 The Data Provider hereby irrevocably authorizes the Data Recipient to:\n(a) migrate the Original Study Data from on-premises infrastructure to the Azure Cloud Platform;\n(b) store and process the Original Study Data and all derived data sets in the Azure Cloud Platform; and\n(c) implement the tiered Data Access Levels system described in this Amendment.\n\n2.20 External Researcher Access Authorization\n\n2.20.1 The Data Provider hereby explicitly and irrevocably authorizes and grants permission for appropriately de-identified data derived from the Original Study Data (Level 2 Data) to be made available to External Researchers who are not members of the HE²AT Center Consortium, subject to the following conditions:\n(a) All External Researcher access requests must be reviewed and approved by the Data Access Committee.\n(b) External Researchers must sign a legally binding Data Use Agreement.\n(c) External Researchers must commit to appropriate citation of both the HE²AT Center and the original data sources.\n(d) External Researchers will only have access to Level 2 data (de-identified).\n(e) All External Researcher access will be monitored and logged.\n(f) The Data Access Committee shall maintain the right to revoke access for any External Researcher.\n\n2.21 Post-Project Data Use and Long-Term Data Retention Authorization\n\n2.21.1 The Data Provider hereby explicitly and irrevocably authorizes and grants permission that the data derived from the Original Study Data shall be retained and may continue to be used beyond the conclusion of the HE²AT Center Project as follows:\n(a) Level 0 Data (Original Study Data): Shall be retained for a period of 5 (five) years.\n(b) Level 1 Data (Consortium Shared Data): Shall be retained for a period of 10 (ten) years.\n(c) Level 2 Data (De-identified Data): Shall be retained indefinitely as a scientific resource.\n(d) Level 3 Data (Inferential Data): Shall be retained indefinitely as an open scientific resource.\n\n2.18 The Data Provider'
        },
        # Add new clause in Section 12
        {
            'type': 'addition',
            'search_text': '12.5 The provisions of this Agreement that by their nature are intended to survive termination or expiration of the Agreement shall survive such termination or expiration and shall remain in full force and effect.',
            'replacement': '\n\n12.6 The Parties expressly acknowledge and agree that the provisions of Sections 2.20 and 2.21 of this Amendment regarding External Researcher Access and Post-Project Data Use shall survive the termination of the Agreement and the conclusion of the HE²AT Center Project.'
        },
        # Renumber clause in Section 12
        {
            'type': 'deletion',
            'search_text': '12.6 The Data Recipient hereby acknowledges',
            'replacement': '12.7 The Data Recipient hereby acknowledges'
        }
    ]
    
    # Also look for the old RP2 document in OLD_DTAs folder
    old_rp2_path = os.path.join("OLD_DTAs", "Final_Formatted_DTA_HEAT002_WHC_RP2_0903 (3).docx")
    if os.path.exists(old_rp2_path):
        old_rp2_output = os.path.join(output_dir, "Old_RP2_DTA_WHC.docx")
        copy_to_destination(old_rp2_path, old_rp2_output)
    
    return add_tracked_changes(input_path, output_path, changes)

def process_rp1_documents():
    """Process the RP1 documents with tracked changes (both UCT and WHC)"""
    # Define base paths
    input_whc_path = os.path.join("OLD_DTAs_Organized", "RP1_DTA_WHC_Main.docx")
    whc_output_dir = os.path.join("Word_Documents_With_Tracked_Changes", "RP1_WHC")
    whc_output_path = os.path.join(whc_output_dir, "RP1_DTA_WHC_Tracked.docx")
    
    # Find UCT document
    uct_input_path = os.path.join("OLD_DTAs_Organized", "Template_DTA_HEAT001_v2.0.docx")
    uct_output_dir = os.path.join("Word_Documents_With_Tracked_Changes", "RP1_UCT") 
    uct_output_path = os.path.join(uct_output_dir, "RP1_DTA_UCT_Tracked.docx")
    
    # Ensure the output directories exist
    os.makedirs(whc_output_dir, exist_ok=True)
    os.makedirs(uct_output_dir, exist_ok=True)
    
    # Copy the old RP1 documents from OLD_DTAs
    old_rp1_whc_path = os.path.join("OLD_DTAs", "WHC_DTA_RP1_0911 (1).docx")
    old_rp1_uct_path = os.path.join("OLD_DTAs", "DTA_Template_HEAT001_v2.0_UCT_DRAFT 10.09.2024_TC.docx")
    
    if os.path.exists(old_rp1_whc_path):
        old_whc_output = os.path.join(whc_output_dir, "Old_RP1_DTA_WHC.docx")
        copy_to_destination(old_rp1_whc_path, old_whc_output)
    
    if os.path.exists(old_rp1_uct_path):
        old_uct_output = os.path.join(uct_output_dir, "Old_RP1_DTA_UCT.docx")
        copy_to_destination(old_rp1_uct_path, old_uct_output)
    
    # Define the changes for RP1 (which will be applied to both UCT and WHC)
    changes = [
        # Section 1 changes - Definitions
        {
            'type': 'addition',
            'search_text': '1.20 "Consortium Shared Data" means',
            'replacement': '\n\n1.21 "Azure Cloud Platform" means the Microsoft Azure cloud computing service that will serve as the HE²AT Center Primary Repository.\n\n1.22 "Cloud Migration" means the process of transferring the Original Study Data and any derived data sets from the existing on-premises infrastructure to the Azure Cloud Platform.\n\n1.23 "Data Access Committee or DAC" means the committee established by the HE²AT Center to review and approve data access requests from external researchers according to established criteria and protocols, which shall continue to function after the conclusion of the HE²AT Center Project.\n\n1.24 "Data Access Levels" means the tiered access system implemented by the HE²AT Center consisting of:\n• Level 0: Original Study Data - Raw, unprocessed data with restricted access to Core Data Team only\n• Level 1: Consortium Shared Data - Processed data shared only among HE²AT Center Consortium partners\n• Level 2: RP1 De-identified Data - Retained by HE²AT Center for approved external researcher access\n• Level 3: Inferential Data - Aggregated and anonymized data available for open access\n\n1.25 "External Researcher" means any qualified researcher who is not a member of the HE²AT Center Consortium but who has been approved by the Data Access Committee to access Level 2 data for specific research purposes.\n\n1.26 "Post-Project Data Use" means the continued storage, access, and use of the data after the conclusion of the HE²AT Center Project in accordance with this Amendment.\n\n1.27 "Successor Governance Entity" means any entity or institution that assumes responsibility for the governance, maintenance, and oversight of the Post-Project Data Repository after the conclusion of the HE²AT Center Project.'
        },
        # Section 2.1 changes
        {
            'type': 'deletion',
            'search_text': 'This Agreement shall commence on the Commencement Date and shall terminate on completion of the HE²AT Project',
            'replacement': 'This Agreement shall commence on the Commencement Date and shall remain in effect in perpetuity with respect to the Post-Project Data Use provisions set forth in Section 2.12 of this Amendment, unless terminated earlier in accordance with the provisions of this Agreement. The Data Provider specifically acknowledges and agrees that the Post-Project Data Use provisions shall survive the termination of the HE²AT Center Project'
        },
        # Add section after 2.4
        {
            'type': 'addition',
            'search_text': '2.4 Data Provider retains ownership of the Original Study Data and retains all rights to distribute the Original Study Data to other third parties.',
            'replacement': '\n\nThe Data Provider hereby grants the Data Recipient a perpetual, irrevocable, worldwide, non-exclusive license to store the Original Study Data in the Azure Cloud Platform as part of the HE²AT Center\'s Cloud Migration.\n\nThe Data Provider hereby authorizes the Data Recipient to process and transform the Original Study Data to create derived data sets at Levels 1, 2, and 3.\n\nThe Data Provider hereby authorizes the Data Recipient to retain the derived data sets for Post-Project Data Use as specifically authorized in this Amendment.\n\nThe Data Provider hereby authorizes the Data Recipient to grant access to Level 2 data to External Researchers in accordance with the procedures set forth in this Amendment.\n\nThis expanded license does not transfer ownership of the Original Study Data, which remains with the Data Provider.'
        },
        # Add text in 2.5
        {
            'type': 'deletion',
            'search_text': 'The Data Provider acknowledges and agrees that initially the Original Study Data shall be accessible only to the Core Team for purposes of',
            'replacement': 'The Data Provider acknowledges and agrees that initially the Original Study Data shall be accessible only to the Core HE²AT Center Data Management Team for purposes of'
        },
        # Add text to end of 2.5
        {
            'type': 'addition',
            'search_text': 'as set out in the HE²AT Center Data Management Plan.',
            'replacement': ' as set out in the HE²AT Center Data Management Plan. Following the Cloud Migration, the Original Study Data will be classified as Level 0 data in the Azure Cloud Platform\'s tiered data access system and will remain accessible only to the Core Data Team.'
        },
        # New sections 2.19-2.21
        {
            'type': 'addition',
            'search_text': '2.18 The Data Provider',
            'replacement': '\n\n2.19 Cloud Storage Infrastructure and Data Migration Authorization\n\n2.19.1 The Data Provider hereby irrevocably authorizes the Data Recipient to:\n(a) migrate the Original Study Data from on-premises infrastructure to the Azure Cloud Platform;\n(b) store and process the Original Study Data and all derived data sets in the Azure Cloud Platform; and\n(c) implement the tiered Data Access Levels system described in this Amendment.\n\n2.20 External Researcher Access Authorization\n\n2.20.1 The Data Provider hereby explicitly and irrevocably authorizes and grants permission for appropriately de-identified data derived from the Original Study Data (Level 2 Data) to be made available to External Researchers who are not members of the HE²AT Center Consortium, subject to the following conditions:\n(a) All External Researcher access requests must be reviewed and approved by the Data Access Committee.\n(b) External Researchers must sign a legally binding Data Use Agreement.\n(c) External Researchers must commit to appropriate citation of both the HE²AT Center and the original data sources.\n(d) External Researchers will only have access to Level 2 data (de-identified).\n(e) All External Researcher access will be monitored and logged.\n(f) The Data Access Committee shall maintain the right to revoke access for any External Researcher.\n\n2.21 Post-Project Data Use and Long-Term Data Retention Authorization\n\n2.21.1 The Data Provider hereby explicitly and irrevocably authorizes and grants permission that the data derived from the Original Study Data shall be retained and may continue to be used beyond the conclusion of the HE²AT Center Project as follows:\n(a) Level 0 Data (Original Study Data): Shall be retained for a period of 5 (five) years.\n(b) Level 1 Data (Consortium Shared Data): Shall be retained for a period of 10 (ten) years.\n(c) Level 2 Data (De-identified Data): Shall be retained indefinitely as a scientific resource.\n(d) Level 3 Data (Inferential Data): Shall be retained indefinitely as an open scientific resource.\n\n2.18 The Data Provider'
        },
        # Add new clause in Section 12
        {
            'type': 'addition',
            'search_text': '12.5 The provisions of this Agreement that by their nature are intended to survive termination or expiration of the Agreement shall survive such termination or expiration and shall remain in full force and effect.',
            'replacement': '\n\n12.6 The Parties expressly acknowledge and agree that the provisions of Sections 2.20 and 2.21 of this Amendment regarding External Researcher Access and Post-Project Data Use shall survive the termination of the Agreement and the conclusion of the HE²AT Center Project.'
        },
        # Renumber clause in Section 12
        {
            'type': 'deletion',
            'search_text': '12.6 The Data Recipient hereby acknowledges',
            'replacement': '12.7 The Data Recipient hereby acknowledges'
        }
    ]
    
    # Process WHC document
    whc_success = add_tracked_changes(input_whc_path, whc_output_path, changes)
    
    # Process UCT document if it exists
    uct_success = False
    if os.path.exists(uct_input_path):
        uct_success = add_tracked_changes(uct_input_path, uct_output_path, changes)
    else:
        print(f"UCT document not found at {uct_input_path}")
    
    return whc_success, uct_success

if __name__ == "__main__":
    print("Starting to organize and process Word documents with tracked changes...")
    
    # Process the RP2 document
    result_rp2 = process_rp2_document()
    if result_rp2:
        print("Successfully processed RP2 document")
    else:
        print("Failed to process RP2 document")
    
    # Process the RP1 documents
    result_rp1_whc, result_rp1_uct = process_rp1_documents()
    if result_rp1_whc:
        print("Successfully processed RP1 WHC document")
    else:
        print("Failed to process RP1 WHC document")
        
    if result_rp1_uct:
        print("Successfully processed RP1 UCT document")
    else:
        print("Failed to process RP1 UCT document")
    
    # Also copy the original files to the appropriate folders
    for src_path, dest_dir, new_name in [
        (os.path.join("OLD_DTAs_Organized", "RP1_DTA_WHC_Main.docx"), 
         os.path.join("Word_Documents_With_Tracked_Changes", "RP1_WHC"), 
         "Original_RP1_DTA_WHC.docx"),
        
        (os.path.join("OLD_DTAs_Organized", "Template_DTA_HEAT001_v2.0.docx"), 
         os.path.join("Word_Documents_With_Tracked_Changes", "RP1_UCT"), 
         "Original_RP1_DTA_UCT.docx"),
        
        (os.path.join("OLD_DTAs_Organized", "RP2_DTA_WHC_Main.docx"), 
         os.path.join("Word_Documents_With_Tracked_Changes", "RP2_WHC"), 
         "Original_RP2_DTA_WHC.docx")
    ]:
        if os.path.exists(src_path):
            dest_path = os.path.join(dest_dir, new_name)
            copy_to_destination(src_path, dest_path)
    
    # Copy the existing tracked changes files to the appropriate folders
    existing_tracked_files = [
        ("/Users/craig/Desktop/Legal-amendment/RP1_DTA_WHC_Tracked.docx", 
         "/Users/craig/Desktop/Legal-amendment/Word_Documents_With_Tracked_Changes/RP1_WHC/"),
        ("/Users/craig/Desktop/Legal-amendment/RP2_DTA_WHC_Tracked.docx", 
         "/Users/craig/Desktop/Legal-amendment/Word_Documents_With_Tracked_Changes/RP2_WHC/")
    ]
    
    for src, dest in existing_tracked_files:
        if os.path.exists(src):
            shutil.copy2(src, dest)
            print(f"Copied {src} to {dest}")
    
    print("Document processing and organization complete!")
