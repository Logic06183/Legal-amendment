"""
Enhanced Legal Styling Converter
Creates elegantly styled legal documents based on the formatting standards in Amendment.tex and Cover letter.tex
"""

import os
import re
import sys
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm
from docx.enum.section import WD_SECTION

def add_page_number(paragraph):
    """Add page numbers to footer"""
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    
    run = paragraph.add_run()
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_header(doc, header_text):
    """Add header with center-aligned text"""
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_run = header_para.add_run(header_text)
    header_run.font.name = 'Times New Roman'
    header_run.font.size = Pt(10)
    header_run.font.bold = True
    
    # Add header line
    header_para.paragraph_format.space_after = Pt(12)
    header_para = header.add_paragraph()
    
    return header

def apply_legal_document_styling(doc):
    """Apply professional legal document styling based on HE2AT templates"""
    # Set margin - standard legal document margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set paragraph spacing
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # Add Header and Footer
    add_header(doc, "HE²AT CENTER")
    
    # Add footer with page numbers
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    add_page_number(footer_para)
    
    return doc

def convert_markdown_to_styled_word(input_file, output_file):
    """
    Convert a markdown file with tracked changes to an elegantly styled Word document.
    
    Args:
        input_file: Path to the markdown file with tracked changes
        output_file: Path to save the styled Word document
    """
    print(f"Converting {input_file} to {output_file} with enhanced legal styling...")
    
    # Read the markdown file
    with open(input_file, 'r', encoding='utf-8') as file:
        content = file.read()
    
    # Create a new Word document
    doc = Document()
    doc = apply_legal_document_styling(doc)
    
    # Extract the title from the first line
    first_line = content.split('\n', 1)[0].strip('# ')
    
    # Add document title with professional formatting
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.add_run(first_line)
    title_run.font.size = Pt(16)
    title_run.bold = True
    title_paragraph.space_after = Pt(16)
    
    # Add a legend with professional formatting
    legend_paragraph = doc.add_paragraph()
    legend_paragraph.paragraph_format.space_before = Pt(6)
    legend_paragraph.paragraph_format.space_after = Pt(12)
    legend_paragraph.add_run("LEGEND: ").bold = True
    
    legend_run1 = legend_paragraph.add_run("Red strikethrough text")
    legend_run1.font.color.rgb = RGBColor(255, 0, 0)
    legend_run1.font.strike = True
    legend_paragraph.add_run(" = deleted text; ")
    
    legend_run2 = legend_paragraph.add_run("Blue text")
    legend_run2.font.color.rgb = RGBColor(0, 0, 255)
    legend_paragraph.add_run(" = added text")
    
    # Add a separator
    hr_paragraph = doc.add_paragraph()
    hr_paragraph.paragraph_format.space_before = Pt(0)
    hr_paragraph.paragraph_format.space_after = Pt(12)
    hr_pFmt = hr_paragraph.paragraph_format
    hr_pFmt.left_indent = -Inches(0.5)
    hr_pFmt.right_indent = -Inches(0.5)
    hr_run = hr_paragraph.add_run()
    hr_rFmt = hr_run.font
    hr_rFmt.color.rgb = RGBColor(0, 0, 0)
    hr_run.add_break()
    
    # Process the markdown content line by line
    lines = content.split('\n')
    in_paragraph = False
    current_text = ""
    skip_lines = 7  # Skip title and legend lines
    
    for i, line in enumerate(lines):
        if i < skip_lines:
            continue
            
        # Check if the line is empty - it marks the end of a paragraph
        if not line.strip():
            if in_paragraph:
                # Process and add the current paragraph
                process_paragraph(doc, current_text)
                current_text = ""
                in_paragraph = False
            continue
        
        # Check for headers
        if line.strip().startswith('#'):
            if in_paragraph:
                process_paragraph(doc, current_text)
                current_text = ""
                in_paragraph = False
            
            # Count header level and extract text
            header_level = 0
            for char in line:
                if char == '#':
                    header_level += 1
                else:
                    break
            
            header_text = line.lstrip('#').strip()
            
            # Create header with appropriate formatting
            header = doc.add_paragraph()
            header.paragraph_format.space_before = Pt(12)
            header.paragraph_format.space_after = Pt(6)
            header.paragraph_format.keep_with_next = True
            
            # Add appropriate styling based on header level
            header_run = header.add_run(header_text)
            header_run.bold = True
            
            if header_level == 1:
                header_run.font.size = Pt(16)
                header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif header_level == 2:
                header_run.font.size = Pt(14)
            else:
                header_run.font.size = Pt(13)
            
            continue
        
        # Start a new paragraph or continue with the current one
        if not in_paragraph:
            in_paragraph = True
            current_text = line
        else:
            current_text += " " + line
    
    # Process the last paragraph if any
    if current_text:
        process_paragraph(doc, current_text)
    
    # Save the document
    doc.save(output_file)
    print(f"Successfully converted with enhanced legal styling: {output_file}")
    return True

def process_paragraph(doc, text):
    """
    Process a single paragraph of text and add it to the document with tracked changes.
    Enhanced for professional legal document styling.
    """
    # Create a new paragraph with proper spacing
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # Process the text to extract normal text, deleted text, and added text
    remaining_text = text
    while remaining_text:
        # Check for blue text (additions)
        blue_match = re.search(r'<span style=\'color: blue\'>(.*?)</span>', remaining_text)
        # Check for red text (deletions)
        red_match = re.search(r'<span style=\'color: red\'>(.*?)</span>', remaining_text)
        # Check for strikethrough text (deletions)
        strike_match = re.search(r'~~(.*?)~~', remaining_text)
        
        # If there's no markup, add the remaining text normally
        if not blue_match and not red_match and not strike_match:
            run = paragraph.add_run(remaining_text)
            break
        
        # Find the first match
        matches = [m for m in [blue_match, red_match, strike_match] if m]
        if not matches:
            run = paragraph.add_run(remaining_text)
            break
            
        matches.sort(key=lambda m: m.start())
        first_match = matches[0]
        
        # Add text before the match
        if first_match.start() > 0:
            run = paragraph.add_run(remaining_text[:first_match.start()])
        
        # Handle the match based on its type
        if first_match == blue_match:
            # Blue text: additions
            addition_text = blue_match.group(1)
            run = paragraph.add_run(addition_text)
            run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
            remaining_text = remaining_text[blue_match.end():]
        elif first_match == red_match:
            # Red text: deletions
            deletion_text = red_match.group(1)
            run = paragraph.add_run(deletion_text)
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            run.font.strike = True
            remaining_text = remaining_text[red_match.end():]
        elif first_match == strike_match:
            # Strikethrough: deletions
            deletion_text = strike_match.group(1)
            run = paragraph.add_run(deletion_text)
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            run.font.strike = True
            remaining_text = remaining_text[strike_match.end():]
    
    return paragraph

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python enhanced_legal_styling.py input_markdown.md output_word.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found.")
        sys.exit(1)
    
    convert_markdown_to_styled_word(input_file, output_file)
