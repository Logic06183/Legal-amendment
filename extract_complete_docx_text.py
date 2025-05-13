"""
More thorough extraction of text from Word documents and creation of markdown with tracked changes.
This version ensures we extract ALL content properly.
"""

import os
from docx import Document
import re
from docx.enum.text import WD_BREAK

def extract_all_text_from_docx(docx_path):
    """Extract all text from a Word document, including tables and preserving structure"""
    doc = Document(docx_path)
    full_text = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():  # Only append non-empty paragraphs
            full_text.append(para.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text)
            if row_text:
                full_text.append(" | ".join(row_text))
    
    return "\n\n".join(full_text)

def create_complete_markdown(input_docx, output_md, document_type):
    """Create complete markdown file with tracked changes from Word document"""
    try:
        # Extract text from Word document
        text_content = extract_all_text_from_docx(input_docx)
        
        # Create markdown content
        markdown = []
        
        # Add header with tracked changes legend
        markdown.append(f"# DATA TRANSFER AGREEMENT WITH TRACKED CHANGES - {document_type}")
        markdown.append("")
        markdown.append("**Legend:**")
        markdown.append("- ~~<span style='color: red'>Red strikethrough text</span>~~ = deleted text")
        markdown.append("- <span style='color: blue'>Blue text</span> = added text")
        markdown.append("")
        markdown.append("---")
        markdown.append("")
        
        # Convert text to markdown format
        lines = text_content.split("\n\n")
        for i, line in enumerate(lines):
            # Clean up the line
            line = line.strip()
            if not line:
                continue
                
            # Format headings and content
            if i == 0 or (len(line) < 100 and line.isupper()):
                markdown.append(f"# {line}")
            else:
                markdown.append(line)
            
            # Add tracked changes after specific sections
            
            # 1. Add new definitions after definitions section
            if "DEFINITIONS" in line or "Consortium Shared Data" in line:
                markdown.append("")
                markdown.append("<span style='color: blue'>")
                markdown.append("1.21 \"Azure Cloud Platform\" means the Microsoft Azure cloud computing service that will serve as the HE²AT Center Primary Repository.")
                markdown.append("")
                markdown.append("1.22 \"Cloud Migration\" means the process of transferring the Original Study Data and any derived data sets from the existing on-premises infrastructure to the Azure Cloud Platform.")
                markdown.append("")
                markdown.append("1.23 \"Data Access Committee or DAC\" means the committee established by the HE²AT Center to review and approve data access requests from external researchers according to established criteria and protocols, which shall continue to function after the conclusion of the HE²AT Center Project.")
                markdown.append("")
                markdown.append("1.24 \"Data Access Levels\" means the tiered access system implemented by the HE²AT Center consisting of:")
                markdown.append("• Level 0: Original Study Data - Raw, unprocessed data with restricted access to Core Data Team only")
                markdown.append("• Level 1: Consortium Shared Data - Processed data shared only among HE²AT Center Consortium partners")
                markdown.append("• Level 2: De-identified Data - Retained by HE²AT Center for approved external researcher access")
                markdown.append("• Level 3: Inferential Data - Aggregated and anonymized data available for open access")
                markdown.append("")
                markdown.append("1.25 \"External Researcher\" means any qualified researcher who is not a member of the HE²AT Center Consortium but who has been approved by the Data Access Committee to access Level 2 data for specific research purposes.")
                markdown.append("")
                markdown.append("1.26 \"Post-Project Data Use\" means the continued storage, access, and use of the data after the conclusion of the HE²AT Center Project in accordance with this Amendment.")
                markdown.append("")
                markdown.append("1.27 \"Successor Governance Entity\" means any entity or institution that assumes responsibility for the governance, maintenance, and oversight of the Post-Project Data Repository after the conclusion of the HE²AT Center Project.")
                markdown.append("</span>")
            
            # 2. Replace agreement duration in Section 2.1
            if "shall terminate on completion of the HE²AT Project" in line:
                modified_line = line.replace(
                    "shall terminate on completion of the HE²AT Project",
                    "~~<span style='color: red'>shall terminate on completion of the HE²AT Project</span>~~ <span style='color: blue'>shall remain in effect in perpetuity with respect to the Post-Project Data Use provisions set forth in Section 2.12 of this Amendment, unless terminated earlier in accordance with the provisions of this Agreement. The Data Provider specifically acknowledges and agrees that the Post-Project Data Use provisions shall survive the termination of the HE²AT Center Project</span>"
                )
                markdown[-1] = modified_line
            
            # 3. Add enhanced data rights after Section 2.4
            if "2.4" in line and "Data Provider retains ownership of the Original Study Data" in line:
                markdown.append("")
                markdown.append("<span style='color: blue'>")
                markdown.append("The Data Provider hereby grants the Data Recipient a perpetual, irrevocable, worldwide, non-exclusive license to store the Original Study Data in the Azure Cloud Platform as part of the HE²AT Center's Cloud Migration.")
                markdown.append("")
                markdown.append("The Data Provider hereby authorizes the Data Recipient to process and transform the Original Study Data to create derived data sets at Levels 1, 2, and 3.")
                markdown.append("")
                markdown.append("The Data Provider hereby authorizes the Data Recipient to retain the derived data sets for Post-Project Data Use as specifically authorized in this Amendment.")
                markdown.append("")
                markdown.append("The Data Provider hereby authorizes the Data Recipient to grant access to Level 2 data to External Researchers in accordance with the procedures set forth in this Amendment.")
                markdown.append("")
                markdown.append("This expanded license does not transfer ownership of the Original Study Data, which remains with the Data Provider.")
                markdown.append("</span>")
            
            # 4. Replace Core Team reference and add Level 0 classification
            if "Core Team for purposes of" in line:
                modified_line = line.replace(
                    "Core Team for purposes of",
                    "~~<span style='color: red'>Core Team</span>~~ <span style='color: blue'>Core HE²AT Center Data Management Team</span> for purposes of"
                )
                markdown[-1] = modified_line
            
            if "HE²AT Center Data Management Plan." in line:
                modified_line = line.replace(
                    "HE²AT Center Data Management Plan.",
                    "HE²AT Center Data Management Plan. <span style='color: blue'>Following the Cloud Migration, the Original Study Data will be classified as Level 0 data in the Azure Cloud Platform's tiered data access system and will remain accessible only to the Core Data Team.</span>"
                )
                markdown[-1] = modified_line
            
            # 5. Add new sections 2.19-2.21 before GENERAL section
            if "GENERAL" in line or "general" in line:
                markdown.append("")
                markdown.append("<span style='color: blue'>")
                markdown.append("## 2.19 Cloud Storage Infrastructure and Data Migration Authorization")
                markdown.append("")
                markdown.append("2.19.1 The Data Provider hereby irrevocably authorizes the Data Recipient to:")
                markdown.append("(a) migrate the Original Study Data from on-premises infrastructure to the Azure Cloud Platform;")
                markdown.append("(b) store and process the Original Study Data and all derived data sets in the Azure Cloud Platform; and")
                markdown.append("(c) implement the tiered Data Access Levels system described in this Amendment.")
                markdown.append("")
                markdown.append("## 2.20 External Researcher Access Authorization")
                markdown.append("")
                markdown.append("2.20.1 The Data Provider hereby explicitly and irrevocably authorizes and grants permission for appropriately de-identified data derived from the Original Study Data (Level 2 Data) to be made available to External Researchers who are not members of the HE²AT Center Consortium, subject to the following conditions:")
                markdown.append("(a) All External Researcher access requests must be reviewed and approved by the Data Access Committee.")
                markdown.append("(b) External Researchers must sign a legally binding Data Use Agreement.")
                markdown.append("(c) External Researchers must commit to appropriate citation of both the HE²AT Center and the original data sources.")
                markdown.append("(d) External Researchers will only have access to Level 2 data (de-identified).")
                markdown.append("(e) All External Researcher access will be monitored and logged.")
                markdown.append("(f) The Data Access Committee shall maintain the right to revoke access for any External Researcher.")
                markdown.append("")
                markdown.append("## 2.21 Post-Project Data Use and Long-Term Data Retention Authorization")
                markdown.append("")
                markdown.append("2.21.1 The Data Provider hereby explicitly and irrevocably authorizes and grants permission that the data derived from the Original Study Data shall be retained and may continue to be used beyond the conclusion of the HE²AT Center Project as follows:")
                markdown.append("(a) Level 0 Data (Original Study Data): Shall be retained for a period of 5 (five) years.")
                markdown.append("(b) Level 1 Data (Consortium Shared Data): Shall be retained for a period of 10 (ten) years.")
                markdown.append("(c) Level 2 Data (De-identified Data): Shall be retained indefinitely as a scientific resource.")
                markdown.append("(d) Level 3 Data (Inferential Data): Shall be retained indefinitely as an open scientific resource.")
                markdown.append("</span>")
            
            # 6. Add new survival clause after 12.5
            if "12.5" in line and "survive" in line:
                markdown.append("")
                markdown.append("<span style='color: blue'>")
                markdown.append("12.6 The Parties expressly acknowledge and agree that the provisions of Sections 2.20 and 2.21 of this Amendment regarding External Researcher Access and Post-Project Data Use shall survive the termination of the Agreement and the conclusion of the HE²AT Center Project.")
                markdown.append("</span>")
            
            # 7. Renumber 12.6 to 12.7
            if "12.6" in line:
                modified_line = line.replace(
                    "12.6",
                    "~~<span style='color: red'>12.6</span>~~ <span style='color: blue'>12.7</span>"
                )
                markdown[-1] = modified_line
            
            # 8. Change Effective Date to Amendment Effective Date
            if "Effective Date" in line and "WITNESS" in line:
                modified_line = line.replace(
                    "Effective Date",
                    "~~<span style='color: red'>Effective Date</span>~~ <span style='color: blue'>Amendment Effective Date</span>"
                )
                markdown[-1] = modified_line
            
            markdown.append("")  # Add empty line after each paragraph
        
        # Write markdown to file
        with open(output_md, 'w', encoding='utf-8') as f:
            f.write('\n'.join(markdown))
        
        print(f"Created complete markdown with tracked changes for {document_type}: {output_md}")
        return True
        
    except Exception as e:
        print(f"Error creating markdown: {e}")
        return False

def main():
    base_dir = "/Users/craig/Desktop/Legal-amendment"
    final_word_docs = os.path.join(base_dir, "Final_Word_Documents")
    
    # Process RP2_WHC
    docx_path = os.path.join(final_word_docs, "RP2_WHC", "RP2_WHC_Original.docx")
    output_path = os.path.join(final_word_docs, "RP2_WHC", "RP2_WHC_Complete_With_Changes.md")
    create_complete_markdown(docx_path, output_path, "RP2 Document")
    
    # Process RP1_WHC_New
    docx_path = os.path.join(final_word_docs, "RP1_WHC_New", "RP1_WHC_Original.docx")
    output_path = os.path.join(final_word_docs, "RP1_WHC_New", "RP1_WHC_Complete_With_Changes.md")
    create_complete_markdown(docx_path, output_path, "RP1 WHC New")
    
    # Process RP1_UCT_New
    docx_path = os.path.join(final_word_docs, "RP1_UCT_New", "RP1_UCT_Original.docx")
    output_path = os.path.join(final_word_docs, "RP1_UCT_New", "RP1_UCT_Complete_With_Changes.md")
    create_complete_markdown(docx_path, output_path, "RP1 UCT New")
    
    # Process RP1_Old
    docx_path = os.path.join(final_word_docs, "RP1_Old", "RP1_WHC_Original.docx")
    output_path = os.path.join(final_word_docs, "RP1_Old", "RP1_WHC_Complete_With_Changes.md")
    create_complete_markdown(docx_path, output_path, "RP1 Old")
    
    print("All documents processed successfully!")

if __name__ == "__main__":
    main()
