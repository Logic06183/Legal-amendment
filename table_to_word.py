"""
Convert the amendment_changes_table.md to a professionally formatted Word document
"""

import os
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT

def create_word_table_from_markdown(input_file, output_file):
    """Convert markdown table to professionally styled Word document"""
    print(f"Converting {input_file} to {output_file}...")
    
    # Read markdown file
    with open(input_file, 'r', encoding='utf-8') as file:
        markdown_content = file.readlines()
    
    # Create a new Word document
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("HE²AT CENTER LEGAL AMENDMENT CHANGES")
    title_run.bold = True
    title_run.font.size = Pt(16)
    
    # Add subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.add_run("Comprehensive Changes Matrix")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(14)
    
    # Skip title line and separator line
    table_content = [line.strip() for line in markdown_content if line.strip()]
    if len(table_content) < 3:
        print("Error: Not enough content for a table")
        return False
    
    # Parse header row to determine number of columns
    header_row = table_content[1].strip('| ').split('|')
    num_columns = len(header_row)
    
    # Create table with appropriate styling
    table = doc.add_table(rows=1, cols=num_columns)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths (first column wider for section names)
    for i, cell in enumerate(table.rows[0].cells):
        if i == 0:  # Section column
            cell.width = Inches(1.5)
        elif i == 1:  # Change Description column
            cell.width = Inches(2.5)
        else:  # Institution columns
            cell.width = Inches(0.8)
    
    # Add header row with styling
    header_cells = table.rows[0].cells
    for i, text in enumerate(header_row):
        header_cells[i].text = text.strip()
        header_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(12)
    
    # Process content rows
    for i in range(3, len(table_content)):
        # Skip separator row
        if '---' in table_content[i]:
            continue
            
        cells = table_content[i].strip('| ').split('|')
        row_cells = table.add_row().cells
        
        for j, text in enumerate(cells):
            cell_text = text.strip()
            
            # Handle bold text in first column (section titles)
            if j == 0 and cell_text.startswith('**') and cell_text.endswith('**'):
                row_cells[j].text = cell_text.strip('**')
                for paragraph in row_cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            else:
                row_cells[j].text = cell_text
            
            # Center alignment for checkmark columns
            if j >= 2:  # Institution columns
                row_cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add spacing after table
    doc.add_paragraph()
    
    # Add footer note
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_run = footer_para.add_run("This table summarizes all amendments to the HE²AT Center Data Transfer Agreements")
    footer_run.italic = True
    
    # Save the Word document
    doc.save(output_file)
    print(f"Successfully created: {output_file}")
    return True

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python table_to_word.py input_table.md output_table.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found.")
        sys.exit(1)
    
    create_word_table_from_markdown(input_file, output_file)
