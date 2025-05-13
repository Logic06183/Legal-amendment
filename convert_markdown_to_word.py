"""
Convert Markdown files with tracked changes (colored text) to Word documents with proper tracked changes.
This script will take the markdown files from Final_Word_Documents and create proper Word documents
that lawyers can easily edit via email.
"""

import os
import re
import sys
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_BREAK

def convert_markdown_to_word(markdown_file, output_docx):
    """
    Convert a markdown file with colored text changes to a Word document with proper tracked changes.
    
    Args:
        markdown_file: Path to markdown file with HTML-style tracked changes
        output_docx: Path to save the Word document with tracked changes
    """
    print(f"Converting {markdown_file} to {output_docx}...")
    
    # Read the markdown file
    with open(markdown_file, 'r', encoding='utf-8') as file:
        content = file.read()
    
    # Create a new Word document
    doc = Document()
    
    # Process the markdown content line by line
    lines = content.split('\n')
    in_paragraph = False
    current_text = ""
    
    for line in lines:
        # Check if the line is empty - it marks the end of a paragraph
        if not line.strip():
            if in_paragraph:
                # Process and add the current paragraph
                process_paragraph(doc, current_text)
                current_text = ""
                in_paragraph = False
            continue
        
        # Start a new paragraph or continue with the current one
        if not in_paragraph:
            in_paragraph = True
            current_text = line
        else:
            current_text += " " + line
    
    # Process the last paragraph if any
    if current_text:
        process_paragraph(doc, current_text)
    
    # Save the document
    doc.save(output_docx)
    print(f"Successfully saved: {output_docx}")
    return True

def process_paragraph(doc, text):
    """
    Process a single paragraph of text and add it to the document with tracked changes.
    """
    # Create a new paragraph
    paragraph = doc.add_paragraph()
    
    # Process the text to extract normal text, deleted text, and added text
    while text:
        # Check for blue text (additions)
        blue_match = re.search(r'<span style=\'color: blue\'>(.*?)</span>', text)
        # Check for red text (deletions)
        red_match = re.search(r'<span style=\'color: red\'>(.*?)</span>', text, re.DOTALL)
        # Check for strikethrough text (deletions)
        strike_match = re.search(r'~~(.*?)~~', text)
        
        # If there's no markup, add the remaining text normally
        if not blue_match and not red_match and not strike_match:
            run = paragraph.add_run(text)
            break
        
        # Find the first match
        matches = [m for m in [blue_match, red_match, strike_match] if m]
        if not matches:
            run = paragraph.add_run(text)
            break
            
        matches.sort(key=lambda m: m.start())
        first_match = matches[0]
        
        # Add text before the match
        if first_match.start() > 0:
            run = paragraph.add_run(text[:first_match.start()])
        
        # Handle the match based on its type
        if first_match == blue_match:
            # Blue text: additions
            addition_text = blue_match.group(1)
            run = paragraph.add_run(addition_text)
            run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
            text = text[blue_match.end():]
        elif first_match == red_match or first_match == strike_match:
            # Red text or strikethrough: deletions
            if first_match == red_match:
                deletion_text = red_match.group(1)
                text = text[red_match.end():]
            else:
                deletion_text = strike_match.group(1)
                text = text[strike_match.end():]
                
            run = paragraph.add_run(deletion_text)
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            run.font.strike = True
    
    # Add paragraph break
    paragraph.add_run().add_break(WD_BREAK.LINE)

def main():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Define input and output paths
    input_dirs = [
        os.path.join(base_dir, "Final_Word_Documents", "RP1_UCT_New"),
        os.path.join(base_dir, "Final_Word_Documents", "RP1_WHC_New"),
        os.path.join(base_dir, "Final_Word_Documents", "RP2_WHC")
    ]
    
    # Create output directory if it doesn't exist
    output_dir = os.path.join(base_dir, "Lawyer_Editable_Word_Docs")
    os.makedirs(output_dir, exist_ok=True)
    
    # Process each directory
    for input_dir in input_dirs:
        dir_name = os.path.basename(input_dir)
        output_subdir = os.path.join(output_dir, dir_name)
        os.makedirs(output_subdir, exist_ok=True)
        
        # Find all markdown files
        for file_name in os.listdir(input_dir):
            if file_name.endswith("_With_Changes.md"):
                markdown_file = os.path.join(input_dir, file_name)
                # Create output file name
                base_name = file_name.replace("_With_Changes.md", "")
                output_file = os.path.join(output_subdir, f"{base_name}_Lawyer_Editable.docx")
                
                # Convert the file
                convert_markdown_to_word(markdown_file, output_file)
    
    print("\nConversion complete! Lawyer-editable Word documents have been created in:")
    print(output_dir)
    print("\nThese documents can be emailed to lawyers for editing.")

if __name__ == "__main__":
    main()
