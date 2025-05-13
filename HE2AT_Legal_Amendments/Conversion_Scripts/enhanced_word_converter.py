"""
Enhanced Markdown to Word converter specifically designed for legal documents with tracked changes.
This script creates properly formatted Word documents that preserve tracked changes in a way
legal professionals can easily review and edit.
"""

import os
import re
import sys
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    
    run = paragraph.add_run()
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def convert_markdown_to_word(markdown_file, output_docx, include_legend=True):
    """
    Convert a markdown file with colored text changes to a Word document with proper legal formatting.
    
    Args:
        markdown_file: Path to markdown file with HTML-style tracked changes
        output_docx: Path to save the Word document with tracked changes
        include_legend: Whether to include a legend explaining the formatting
    """
    print(f"Converting {markdown_file} to {output_docx}...")
    
    # Read the markdown file
    with open(markdown_file, 'r', encoding='utf-8') as file:
        content = file.read()
    
    # Create a new Word document with legal formatting
    doc = Document()
    
    # Set font for entire document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Add title from the first line of the markdown
    first_line = content.split('\n', 1)[0].strip('# ')
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(first_line)
    title_run.bold = True
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()  # Add some space
    
    # Add a legend if requested
    if include_legend:
        legend_paragraph = doc.add_paragraph()
        legend_paragraph.add_run("LEGEND: ").bold = True
        legend_run1 = legend_paragraph.add_run("Red strikethrough text")
        legend_run1.font.color.rgb = RGBColor(255, 0, 0)
        legend_run1.font.strike = True
        legend_paragraph.add_run(" = deleted text; ")
        legend_run2 = legend_paragraph.add_run("Blue text")
        legend_run2.font.color.rgb = RGBColor(0, 0, 255)
        legend_paragraph.add_run(" = added text")
        doc.add_paragraph("---")
        doc.add_paragraph()  # Add some space
    
    # Process the markdown content line by line
    lines = content.split('\n')
    in_paragraph = False
    current_text = ""
    skip_lines = 7 if include_legend else 1  # Skip title and legend
    
    for i, line in enumerate(lines):
        if i < skip_lines:
            continue
            
        # Check if the line is empty - it marks the end of a paragraph
        if not line.strip():
            if in_paragraph:
                # Process and add the current paragraph
                process_paragraph(doc, current_text)
                current_text = ""
                in_paragraph = False
            continue
        
        # Check for headers
        if line.startswith('#'):
            if in_paragraph:
                process_paragraph(doc, current_text)
                current_text = ""
                in_paragraph = False
            
            header_level = len(line.split(' ')[0])
            header_text = line.lstrip('#').strip()
            
            header = doc.add_paragraph()
            header_run = header.add_run(header_text)
            header_run.bold = True
            if header_level == 1:
                header_run.font.size = Pt(16)
            elif header_level == 2:
                header_run.font.size = Pt(14)
            else:
                header_run.font.size = Pt(13)
            continue
        
        # Start a new paragraph or continue with the current one
        if not in_paragraph:
            in_paragraph = True
            current_text = line
        else:
            current_text += " " + line
    
    # Process the last paragraph if any
    if current_text:
        process_paragraph(doc, current_text)
    
    # Add page numbers in footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    add_page_number(footer_para)
    
    # Save the document
    doc.save(output_docx)
    print(f"Successfully saved: {output_docx}")
    return True

def process_paragraph(doc, text):
    """
    Process a single paragraph of text and add it to the document with tracked changes.
    Specifically optimized for legal document formatting.
    """
    # Create a new paragraph
    paragraph = doc.add_paragraph()
    
    # Process the text to extract normal text, deleted text, and added text
    while text:
        # Check for blue text (additions)
        blue_match = re.search(r'<span style=\'color: blue\'>(.*?)</span>', text)
        # Check for red text (deletions)
        red_match = re.search(r'<span style=\'color: red\'>(.*?)</span>', text, re.DOTALL)
        # Check for strikethrough text (deletions)
        strike_match = re.search(r'~~(.*?)~~', text)
        
        # If there's no markup, add the remaining text normally
        if not blue_match and not red_match and not strike_match:
            run = paragraph.add_run(text)
            break
        
        # Find the first match
        matches = [m for m in [blue_match, red_match, strike_match] if m]
        if not matches:
            run = paragraph.add_run(text)
            break
            
        matches.sort(key=lambda m: m.start())
        first_match = matches[0]
        
        # Add text before the match
        if first_match.start() > 0:
            run = paragraph.add_run(text[:first_match.start()])
        
        # Handle the match based on its type
        if first_match == blue_match:
            # Blue text: additions
            addition_text = blue_match.group(1)
            run = paragraph.add_run(addition_text)
            run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
            text = text[blue_match.end():]
        elif first_match == red_match or first_match == strike_match:
            # Red text or strikethrough: deletions
            if first_match == red_match:
                deletion_text = red_match.group(1)
                text = text[red_match.end():]
            else:
                deletion_text = strike_match.group(1)
                text = text[strike_match.end():]
                
            run = paragraph.add_run(deletion_text)
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            run.font.strike = True
    
    return paragraph

def main():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Define input and output paths
    input_dirs = [
        os.path.join(base_dir, "Final_Word_Documents", "RP1_UCT_New"),
        os.path.join(base_dir, "Final_Word_Documents", "RP1_WHC_New"),
        os.path.join(base_dir, "Final_Word_Documents", "RP2_WHC"),
        os.path.join(base_dir, "Final_Word_Documents", "RP1_Old")
    ]
    
    # Create output directory if it doesn't exist
    output_dir = os.path.join(base_dir, "Legal_Review_Word_Documents")
    os.makedirs(output_dir, exist_ok=True)
    
    # Process each directory
    for input_dir in input_dirs:
        dir_name = os.path.basename(input_dir)
        output_subdir = os.path.join(output_dir, dir_name)
        os.makedirs(output_subdir, exist_ok=True)
        
        # Find all markdown files
        for file_name in os.listdir(input_dir):
            if file_name.endswith("_With_Changes.md"):
                markdown_file = os.path.join(input_dir, file_name)
                # Create output file name
                base_name = file_name.replace("_With_Changes.md", "")
                output_file = os.path.join(output_subdir, f"{base_name}_Legal_Review.docx")
                
                # Convert the file
                convert_markdown_to_word(markdown_file, output_file)
    
    print("\nConversion complete! Legal review Word documents have been created in:")
    print(output_dir)
    print("\nThese documents include:")
    print("1. Proper formatting for legal review")
    print("2. Red strikethrough text for deletions")
    print("3. Blue text for additions")
    print("4. Page numbers in the footer")
    print("5. Proper spacing and header formatting")
    print("\nLawyers can now edit these documents directly in Word.")

if __name__ == "__main__":
    main()
